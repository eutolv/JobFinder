#!/usr/bin/env python3
# scrapperL.py — LinkedIn Scraper V4.1 (Tolv Edition)
# Totalmente reescrito com:
# - Normalização de URLs
# - Filtro de experiência ajustável
# - Retry inteligente
# - Barras de progresso (global + página)
# - JSON + DOCX em pasta organizada
# - Log limpo

import re
import time
import json
import logging
import os
from datetime import datetime
from urllib.parse import urljoin, urlparse, urlunparse
from typing import List, Dict, Any, Optional

from playwright.sync_api import sync_playwright, Page, TimeoutError as PlaywrightTimeoutError
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ---------------------------------------------------------
# LOGGING
# ---------------------------------------------------------
logging.basicConfig(
    format="%(asctime)s [%(levelname)s] %(message)s",
    level=logging.INFO
)
logger = logging.getLogger("linkedin_v4")

# ---------------------------------------------------------
# CONFIG
# ---------------------------------------------------------
ALLOW_BRAZIL = True
ALLOW_UNSPECIFIED_EXPERIENCE = True
MAX_RETRIES = 3

REGEX_COUNTRIES = re.compile(
    r"\b(" +
    r"|".join([
        "united states", "usa", "canada", "uk", "united kingdom", "india",
        "poland", "ukraine", "philippines", "australia", "new zealand",
        "germany", "france", "spain", "italy", "mexico", "argentina",
        "colombia", "peru", "chile", "ireland", "netherlands", "belgium",
        "sweden", "norway", "denmark", "finland", "japan", "china",
        "singapore", "south africa", "nigeria", "russia"
    ]) +
    r")\b", re.I
)

REGEX_CITIES = re.compile(
    r"\b(" +
    r"|".join([
        "são paulo", "sao paulo", "rio de janeiro", "belo horizonte",
        "curitiba", "porto alegre", "florianopolis", "brasilia",
        "campinas", "recife", "salvador", "fortaleza",
        "toronto", "vancouver", "new york", "london", "paris",
        "berlin", "warsaw", "madrid"
    ]) +
    r")\b", re.I
)

REGEX_TIMEZONES = re.compile(r"\b(utc|gmt|est|pst|cet|bst|aest|gmt[+-]?\d+|utc[+-]?\d+)\b", re.I)

DISALLOWED_LEVEL = re.compile(
    r"\b(senior|sênior|sr\b|sr\.|lead|principal|architect|manager|director|coordinator|supervisor|head of)\b",
    re.I
)

EXPERIENCE_ALLOWED = [
    re.compile(p, re.I) for p in [
        r"\b0\s*(?:years|yrs|anos)?\b",
        r"\bno experience\b",
        r"\bno previous experience\b",
        r"\bentry[-\s]?level\b",
        r"\bjunior\b",
        r"\btrainee\b",
        r"\bgraduate\b",
        r"\bestagi",
        r"\b0[–\-]\s*1\s*years?\b",
        r"\bup to 1 year\b",
        r"\b1\s*year\b",
        r"\b1\s*ano\b",
        r"\b0-1\b"
    ]
]

EXPERIENCE_BLOCK = re.compile(
    r"\b([2-9]\+?\s*(?:years|yrs|anos?)|[2-9]\s*years?)\b",
    re.I
)

ALLOWED_AREAS = [
    "it support", "technical support", "help desk", "support engineer",
    "service desk", "desktop support",
    "soc ", "soc analyst", "security operations", "cybersecurity analyst",
    "incident response", "blue team", "threat analyst",
    "sysadmin", "linux admin", "windows admin",
    "network analyst", "noc analyst", "monitoring analyst",
    "infrastructure analyst",
    "devops junior", "sre junior", "automation engineer"
]
ALLOWED_AREAS_RE = re.compile("|".join(re.escape(a) for a in ALLOWED_AREAS), re.I)

SKILL_BONUS = [
    "splunk", "siem", "python", "linux", "ubuntu", "kali",
    "active directory", "windows server", "tcp/ip", "firewall",
    "vpn", "wireshark", "logs", "phishing", "malware",
    "nmap", "burp", "metasploit", "endpoint"
]

URGENT_TERMS = re.compile(
    r"\b(" + r"|".join([
        "urgent", "urgently", "urgent hire",
        "start asap", "start immediately", "immediate start",
        "hiring now", "contratação imediata", "início imediato"
    ]) + r")\b", re.I
)

SPAM_PATTERNS = re.compile(
    r"\b(template|sample job|test job|mock job|how to apply|demonstration|example job)\b",
    re.I
)

# ---------------------------------------------------------
# HELPERS
# ---------------------------------------------------------
def clean_text(text: Optional[str]) -> str:
    if not text:
        return ""
    return re.sub(r'\s+', ' ', text).strip()

def normalize_job_url(href: str) -> str:
    try:
        parsed = urlparse(href)
        clean = urlunparse((parsed.scheme, parsed.netloc, parsed.path.rstrip("/"), "", "", ""))
        return clean
    except:
        return href

def matches_area(text: str) -> bool:
    return bool(ALLOWED_AREAS_RE.search(text))

def matches_experience(text: str) -> bool:
    t = text.lower()

    if EXPERIENCE_BLOCK.search(t):
        return False

    for p in EXPERIENCE_ALLOWED:
        if p.search(t):
            return True

    return ALLOW_UNSPECIFIED_EXPERIENCE

def score_skills(text: str) -> int:
    t = text.lower()
    return sum(1 for s in SKILL_BONUS if s in t)

def is_spam(t): return bool(SPAM_PATTERNS.search(t))
def is_senior(t): return bool(DISALLOWED_LEVEL.search(t))
def is_urgent(t): return bool(URGENT_TERMS.search(t))

def is_worldwide_ok(text: str) -> bool:
    t = text.lower()

    if "brazil" in t or "brasil" in t:
        return True

    if REGEX_COUNTRIES.search(t):
        return False

    if REGEX_CITIES.search(t):
        city = REGEX_CITIES.search(t).group(0).lower()
        br = ["são paulo", "sao paulo", "rio de janeiro", "belo horizonte",
              "curitiba", "porto alegre", "brasilia"]
        if any(c in city for c in br) and ALLOW_BRAZIL:
            return True
        return False

    if REGEX_TIMEZONES.search(t):
        return False

    blockers = [
        "must be in", "must reside", "must live", "based in", "located in",
        "only", "not hiring outside", "eligible to work in",
        "work authorization"
    ]
    for b in blockers:
        if b in t:
            return False

    return True

# ---------------------------------------------------------
# PAGE EXTRACTOR
# ---------------------------------------------------------
def get_text(page: Page, selectors: List[str]) -> str:
    for sel in selectors:
        try:
            el = page.query_selector(sel)
            if el:
                txt = el.inner_text(timeout=2000)
                if txt:
                    return clean_text(txt)
        except:
            continue
    return ""

def extract_job(page: Page, link: str) -> Optional[Dict[str, Any]]:
    title = get_text(page, ["h1", "h1.jobs-unified-top-card__job-title"])
    company = get_text(page, [".topcard__flavor", ".jobs-unified-top-card__company-name"])
    location = get_text(page, [".jobs-unified-top-card__bullet"])
    description = get_text(page, [
        ".show-more-less-html__markup",
        ".jobs-description-content__text"
    ])

    full = " ".join([title, company, location, description]).lower()

    if is_spam(full): return None
    if is_senior(full): return None
    if not matches_area(full): return None
    if not is_worldwide_ok(full): return None
    if not matches_experience(full): return None

    return {
        "title": title,
        "company": company,
        "location": location or "Remote",
        "description": description,
        "link": link,
        "skills_score": score_skills(full),
        "urgent": is_urgent(full),
        "raw": full
    }

# ---------------------------------------------------------
# PROGRESS BARS
# ---------------------------------------------------------
def progress_bar(prefix: str, step: int, total: int):
    width = 28
    filled = int((step / total) * width)
    bar = "█" * filled + "-" * (width - filled)
    print(f"\r{prefix} [{bar}] {int((step/total)*100)}%", end="", flush=True)

# ---------------------------------------------------------
# SCRAPER
# ---------------------------------------------------------
def scrape(query="IT Support OR SOC", pages=4) -> List[Dict[str, Any]]:
    results = []

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.set_default_timeout(60000)

        base = "https://www.linkedin.com/jobs/search/?keywords="
        encoded = query.replace(" ", "%20")
        full_base = base + encoded + "&location=Worldwide&f_WT=2"

        logger.info("🚀 Iniciando LinkedIn Scraper V4.1 — Tolv Edition\n")

        for pg in range(pages):
            print()  
            logger.info(f"📄 Página {pg+1} de {pages}")

            url = f"{full_base}&start={pg*25}"

            # Retry inteligente
            for attempt in range(1, MAX_RETRIES+1):
                try:
                    page.goto(url)
                    time.sleep(2)
                    break
                except:
                    if attempt == MAX_RETRIES:
                        logger.error(f"Falha ao carregar página {pg+1}")
                        continue
                    time.sleep(attempt * 2)

            anchors = page.query_selector_all("a[href*='/jobs/view/']")
            raw_links = []

            for a in anchors:
                href = a.get_attribute("href")
                if not href:
                    continue
                if href.startswith("/"):
                    href = urljoin("https://www.linkedin.com", href)
                raw_links.append(href)

            raw_count = len(raw_links)

            # Normalização
            seen = set()
            unique_links = []
            for href in raw_links:
                norm = normalize_job_url(href)
                if norm not in seen:
                    seen.add(norm)
                    unique_links.append(href)

            logger.info(f"  → {raw_count} links brutos encontrados")
            logger.info(f"  → {len(unique_links)} links únicos (normalizados)")

            accepted = 0

            # Barra de progresso da página
            for i, link in enumerate(unique_links, 1):
                progress_bar(f"   Página {pg+1}/{pages}", i, len(unique_links))

                try:
                    page.goto(link)
                    time.sleep(0.8)
                except:
                    continue

                job = extract_job(page, link)
                if job:
                    results.append(job)
                    accepted += 1

            print()
            logger.info(f"  → {accepted} vagas aprovadas nesta página")

        browser.close()

    return results

# ---------------------------------------------------------
# EXPORTS
# ---------------------------------------------------------
def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)

def export_docx(jobs: List[Dict[str, Any]], filename: str):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'

    h = doc.add_heading("🌍 LinkedIn Vagas — Tolv Ultimate V4.1", level=1)
    h.runs[0].font.size = Pt(18)
    h.runs[0].font.color.rgb = RGBColor(5, 99, 193)

    meta = doc.add_paragraph()
    meta.add_run("Gerado em: ").bold = True
    meta.add_run(datetime.now().strftime("%d/%m/%Y %H:%M:%S"))
    meta.add_run("\nTotal de vagas: ").bold = True
    meta.add_run(str(len(jobs)))

    urgent = [j for j in jobs if j.get("urgent")]
    high = [j for j in jobs if j.get("skills_score", 0) >= 2]

    if urgent:
        meta.add_run(f"\n🔥 Urgentes: {len(urgent)}")
    if high:
        meta.add_run(f"\n⭐ Alto match: {len(high)}")

    # Prioridades
    doc.add_paragraph("\n" + "="*60)
    doc.add_heading("🔥⭐ PRIORIDADE", level=2)

    prios = sorted(
        [j for j in jobs if j.get("urgent") or j.get("skills_score", 0) >= 2],
        key=lambda x: (-int(x.get("urgent")), -x.get("skills_score"))
    )

    for i, job in enumerate(prios, 1):
        p = doc.add_paragraph()
        prefix = "🔥" if job["urgent"] else ""
        stars = "⭐" * min(job["skills_score"], 3)

        p.add_run(f"{prefix}{stars} {i}. ").bold = True
        add_hyperlink(p, job["title"], job["link"])

        d = doc.add_paragraph()
        info = d.add_run(
            f"🏢 {job['company']}  |  📍 {job['location']}  |  ⭐ Score: {job['skills_score']}"
        )
        info.font.size = Pt(9)
        info.font.color.rgb = RGBColor(110, 110, 110)

    # Outras vagas
    doc.add_paragraph("\n" + "="*60)
    doc.add_heading("📋 OUTRAS VAGAS", level=2)

    other = [j for j in jobs if j not in prios]

    for i, job in enumerate(other, 1):
        p = doc.add_paragraph()
        stars = "⭐" * min(job["skills_score"], 3)
        p.add_run(f"{i}. {stars} ").bold = True
        add_hyperlink(p, job["title"], job["link"])

        d = doc.add_paragraph()
        info = d.add_run(
            f"🏢 {job['company']}  |  📍 {job['location']}  |  ⭐ Score: {job['skills_score']}"
        )
        info.font.size = Pt(9)
        info.font.color.rgb = RGBColor(110, 110, 110)

    doc.save(filename)

# ---------------------------------------------------------
# MAIN
# ---------------------------------------------------------
def main():
    logger.info("🚀 Scraping iniciado (V4.1 — Tolv Edition)")

    jobs = scrape(
        query="IT Support OR SOC OR Infrastructure OR DevOps",
        pages=4
    )

    logger.info(f"✅ Total final de vagas aprovadas: {len(jobs)}")

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    folder = f"results_{timestamp}"
    os.makedirs(folder, exist_ok=True)

    json_path = f"{folder}/linkedin_results.json"
    docx_path = f"{folder}/linkedin_results.docx"

    with open(json_path, "w", encoding="utf8") as f:
        json.dump(jobs, f, indent=2, ensure_ascii=False)

    export_docx(jobs, docx_path)

    logger.info(f"✅ JSON salvo em: {json_path}")
    logger.info(f"✅ DOCX salvo em: {docx_path}")
    logger.info("✅ Finalizado com sucesso!")

if __name__ == "__main__":
    main()
