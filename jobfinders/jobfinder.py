#!/usr/bin/env python3
# jobfinder_minimal.py
# Versão C — Minimalista/otimizada para IT Support (entry-level / junior)
# Sites: Remotive, RemoteOK, WeWorkRemotely, Remote.co, JustRemote, LinkedIn (busca pública)

import re
import time
import sys
from datetime import datetime
from urllib.parse import urljoin, urlparse, urlunparse, parse_qs

import requests
from bs4 import BeautifulSoup
from concurrent.futures import ThreadPoolExecutor, as_completed
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ===========================
# FILTROS
# ===========================
AREA_FILTERS = [
    "it support", "technical support", "tech support", "service desk",
    "help desk", "desktop support", "support engineer", "support technician",
    "support analyst", "infrastructure support", "infrastructure technician",
    "linux support", "sysadmin", "system administrator", "noc",
    "network operations center", "soc", "security operations center",
    "monitoring analyst", "technical support specialist"
]

EXCLUDE_TITLE_TERMS = [
    "senior", "sr ", "sr.", "lead", "manager", "director", "head",
    "principal", "architect", "staff", "chief", "expert", "guru"
]

# Se a vaga explícita pede N anos >= 3, descartamos (gera ruído pra entry)
EXCLUDE_EXPERIENCE_YEARS = 3

# ===========================
# SITES MINIMALISTAS
# ===========================
SITES = {
    "Remotive": "https://remotive.com/remote-jobs",
    "RemoteOK": "https://remoteok.com/remote-jobs",
    "WeWorkRemotely": "https://weworkremotely.com/categories/remote-customer-support-jobs",
    "RemoteCo": "https://remote.co/remote-jobs/customer-service/",
    "JustRemote": "https://justremote.co/remote-jobs?category=Customer%20Support",  # categoria de suporte
    # LinkedIn via busca pública (primeira página)
    "LinkedIn": "https://www.linkedin.com/jobs/search/?" 
}

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
}

MAX_WORKERS = 10
REQUEST_TIMEOUT = 18
SLEEP_BETWEEN_REQUESTS = 0.12

# ===========================
# SESSÃO COM RETRIES
# ===========================
session = requests.Session()
retries = Retry(total=3, backoff_factor=0.5, status_forcelist=(429, 500, 502, 503, 504))
adapter = HTTPAdapter(max_retries=retries)
session.mount("https://", adapter)
session.mount("http://", adapter)
session.headers.update(HEADERS)

# ===========================
# UTILITÁRIOS
# ===========================
def normalize_url(u: str) -> str:
    try:
        p = urlparse(u)
        qs = parse_qs(p.query, keep_blank_values=True)
        qs = {k: v for k, v in qs.items() if not k.lower().startswith("utm") and k.lower() != "fbclid"}
        query = "&".join(f"{k}={v[0]}" for k, v in qs.items())
        return urlunparse((p.scheme, p.netloc, p.path.rstrip("/"), "", query, ""))
    except Exception:
        return u

def fetch_text(url):
    try:
        r = session.get(url, timeout=REQUEST_TIMEOUT)
        r.raise_for_status()
        return r.text
    except Exception:
        return None

def extract_description(soup):
    # meta description
    meta = soup.find("meta", {"name": "description"})
    if meta and meta.get("content"):
        return meta.get("content").strip()
    # article
    article = soup.find("article")
    if article:
        txt = article.get_text(" ").strip()
        if len(txt) > 40:
            return txt
    # common divs
    candidates = soup.find_all(["div","section"], attrs={"class": re.compile(r"(description|job|posting|content|details)", re.I)})
    for c in candidates:
        txt = c.get_text(" ").strip()
        if len(txt) > 40:
            return txt
    # fallback body
    if soup.body:
        return soup.body.get_text(" ").strip()
    return ""

def title_or_h1(soup):
    if soup.title and soup.title.string:
        return soup.title.string.strip()
    h1 = soup.find("h1")
    if h1:
        return h1.get_text(" ").strip()
    return ""

def area_match(text: str) -> bool:
    t = (text or "").lower()
    return any(a in t for a in AREA_FILTERS)

def exclude_title_found(text: str) -> bool:
    t = (text or "").lower()
    return any(e in t for e in EXCLUDE_TITLE_TERMS)

def has_high_experience(text: str) -> bool:
    # procura por "3+ years", "4 years", "5+ years" etc.
    m = re.search(r"(\d+)\+?\s*(years|yrs|anos)", (text or "").lower())
    if m:
        try:
            years = int(m.group(1))
            if years >= EXCLUDE_EXPERIENCE_YEARS:
                return True
        except:
            pass
    return False

def job_passes_filters(title: str, description: str) -> bool:
    combined = f"{title} {description}".lower()
    # REJEITA se título contém termos de senioridade
    if exclude_title_found(title):
        return False
    # REJEITA se experiência pedida >= 3 anos
    if has_high_experience(combined):
        return False
    # PRECISA ser da área (support/it/sysadmin/etc)
    if not area_match(combined):
        return False
    # PASS
    return True

def is_probable_job_url(u: str) -> bool:
    if not u:
        return False
    u = u.lower()
    if any(u.endswith(ext) for ext in [".pdf", ".png", ".jpg", ".jpeg", ".svg", ".zip", ".csv"]):
        return False
    if any(x in u for x in ["#", "mailto:", "javascript:", "tel:"]):
        return False
    # rejeitar claramente páginas de categoria sem slug específico
    if any(x in u for x in ["/categories/", "/jobs/", "/tag/", "/remote-jobs/"]):
        # se tiver slug longo no final, pode ser vaga
        path = urlparse(u).path
        last = path.rstrip("/").split("/")[-1]
        if "-" in last and len(last) > 8:
            return True
        return False
    # heurística por slug com hyphen longo
    path = urlparse(u).path
    last = path.rstrip("/").split("/")[-1]
    if "-" in last and len(last) > 10:
        return True
    return True

# ===========================
# SCRAPERS (Sites minimalistas)
# ===========================

def scrape_remotive():
    print("[+] Remotive")
    url = "https://remotive.com/remote-jobs/search?search=support"
    html = fetch_text(url)
    jobs = []
    if not html:
        return jobs
    soup = BeautifulSoup(html, "html.parser")
    for a in soup.select("a.job-tile-title"):
        title = a.get_text(strip=True)
        href = a.get("href")
        link = urljoin("https://remotive.com", href) if href else None
        if not link:
            continue
        # Remotive list view often doesn't include full description; we'll follow link
        if not is_probable_job_url(link):
            continue
        page = fetch_text(link)
        if not page:
            continue
        psoup = BeautifulSoup(page, "html.parser")
        desc = extract_description(psoup)
        t = title if title else title_or_h1(psoup)
        if job_passes_filters(t, desc):
            jobs.append((t, link))
        time.sleep(SLEEP_BETWEEN_REQUESTS)
    print(f"    → {len(jobs)} vagas filtradas")
    return jobs

def scrape_remoteok():
    print("[+] RemoteOK")
    url = "https://remoteok.com/remote-support-jobs"
    html = fetch_text(url)
    jobs = []
    if not html:
        return jobs
    soup = BeautifulSoup(html, "html.parser")
    for row in soup.select("tr.job"):
        title_el = row.select_one("td.position h2")
        link_el = row.select_one("a.preventLink")
        if not title_el or not link_el:
            continue
        title = title_el.get_text(strip=True)
        href = link_el.get("href")
        link = urljoin("https://remoteok.com", href)
        # follow job page because snippet may be short
        page = fetch_text(link)
        if not page:
            continue
        psoup = BeautifulSoup(page, "html.parser")
        desc = extract_description(psoup)
        t = title if title else title_or_h1(psoup)
        if job_passes_filters(t, desc):
            jobs.append((t, link))
        time.sleep(SLEEP_BETWEEN_REQUESTS)
    print(f"    → {len(jobs)} vagas filtradas")
    return jobs

def scrape_wwr():
    print("[+] WeWorkRemotely")
    url = "https://weworkremotely.com/categories/remote-customer-support-jobs"
    html = fetch_text(url)
    jobs = []
    if not html:
        return jobs
    soup = BeautifulSoup(html, "html.parser")
    for li in soup.select("section.jobs li"):
        a = li.find("a", recursive=False)
        if not a:
            continue
        title = a.get_text(" ").strip()
        href = a.get("href")
        link = urljoin("https://weworkremotely.com", href)
        # follow page
        page = fetch_text(link)
        if not page:
            continue
        psoup = BeautifulSoup(page, "html.parser")
        desc = extract_description(psoup)
        t = title if title else title_or_h1(psoup)
        if job_passes_filters(t, desc):
            jobs.append((t, link))
        time.sleep(SLEEP_BETWEEN_REQUESTS)
    print(f"    → {len(jobs)} vagas filtradas")
    return jobs

def scrape_remoteco():
    print("[+] Remote.co")
    url = "https://remote.co/remote-jobs/customer-service/"
    html = fetch_text(url)
    jobs = []
    if not html:
        return jobs
    soup = BeautifulSoup(html, "html.parser")
    for card in soup.select("div.job-listing a"):
        href = card.get("href")
        if not href:
            continue
        link = urljoin("https://remote.co", href)
        page = fetch_text(link)
        if not page:
            continue
        psoup = BeautifulSoup(page, "html.parser")
        title = title_or_h1(psoup)
        desc = extract_description(psoup)
        if job_passes_filters(title, desc):
            jobs.append((title, link))
        time.sleep(SLEEP_BETWEEN_REQUESTS)
    print(f"    → {len(jobs)} vagas filtradas")
    return jobs

def scrape_justremote():
    print("[+] JustRemote")
    url = "https://justremote.co/remote-jobs?category=Customer%20Support"
    html = fetch_text(url)
    jobs = []
    if not html:
        return jobs
    soup = BeautifulSoup(html, "html.parser")
    for a in soup.select("a.job-link"):
        title = a.get_text(" ").strip()
        href = a.get("href")
        link = urljoin("https://justremote.co", href)
        page = fetch_text(link)
        if not page:
            continue
        psoup = BeautifulSoup(page, "html.parser")
        desc = extract_description(psoup)
        if job_passes_filters(title, desc):
            jobs.append((title, link))
        time.sleep(SLEEP_BETWEEN_REQUESTS)
    print(f"    → {len(jobs)} vagas filtradas")
    return jobs

def scrape_linkedin():
    print("[+] LinkedIn (busca pública)")
    base = "https://www.linkedin.com/jobs/search/?"
    params = {
        "keywords": "IT Support OR Help Desk OR Service Desk OR Technical Support",
        "location": "Worldwide",
        "f_WT": "2",       # remote
        "f_E": "1,2",      # entry level, associate
        "f_TPR": "r604800",# last 7 days
        "f_AL": "true",    # candidatura simplificada
        "position": "1",
        "pageNum": "0"
    }
    url = base + "&" + "&".join(f"{k}={requests.utils.quote(str(v))}" for k,v in params.items())
    html = fetch_text(url)
    jobs = []
    if not html:
        return jobs
    soup = BeautifulSoup(html, "html.parser")
    # LinkedIn public results: anchor with class base-card__full-link
    for a in soup.select("a.base-card__full-link"):
        title = a.get_text(" ").strip()
        href = a.get("href")
        if not href:
            continue
        link = href.split("?")[0]
        # LinkedIn description often not accessible; use title and heuristics
        if job_passes_filters(title, ""):
            jobs.append((title, link))
    print(f"    → {len(jobs)} vagas filtradas")
    return jobs

# ===========================
# DOCX HYPERLINK
# ===========================
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u'); u.set(qn('w:val'), "single"); rPr.append(u)
    color = OxmlElement('w:color'); color.set(qn('w:val'), "0000FF"); rPr.append(color)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# ===========================
# ORQUESTRAÇÃO
# ===========================
def buscar_todos():
    scrapers = [
        scrape_remotive,
        scrape_remoteok,
        scrape_wwr,
        scrape_remoteco,
        scrape_justremote,
        scrape_linkedin
    ]
    all_jobs = []
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as ex:
        futures = {ex.submit(func): func.__name__ for func in scrapers}
        for fut in as_completed(futures):
            try:
                res = fut.result()
                if res:
                    all_jobs.extend(res)
            except Exception as e:
                print(f"[!] Erro em {futures[fut]}: {e}")
    # dedupe por URL normalizada
    uniques = {}
    for t, l in all_jobs:
        n = normalize_url(l)
        if n not in uniques:
            uniques[n] = t
    return uniques

# ===========================
# MAIN
# ===========================
if __name__ == "__main__":
    print("=== JobFinder Minimal (IT Support - Entry) ===")
    start = datetime.now()
    try:
        resultados = buscar_todos()
    except KeyboardInterrupt:
        sys.exit(1)

    print(f"\nTotal final: {len(resultados)} vagas filtradas\n")

    if resultados:
        doc = Document()
        doc.add_heading("Vagas IT Support (Entry-Level) - Minimal", level=1)
        doc.add_paragraph(f"Gerado em: {datetime.now()}\n")
        for url, title in resultados.items():
            p = doc.add_paragraph()
            add_hyperlink(p, title, url)
            doc.add_paragraph("")
        filename = f"it_support_minimal_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"
        doc.save(filename)
        print(f"Arquivo salvo: {filename}")
    else:
        print("Nenhuma vaga encontrada com os filtros definidos.")

    print("Tempo total:", datetime.now() - start)
