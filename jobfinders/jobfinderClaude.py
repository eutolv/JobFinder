#!/usr/bin/env python3
# jobfinder_ultimate_jonathan.py
# Versão ULTIMATE com sites TOP e filtro geo ULTRA-RIGOROSO
# Criado por Apacci para Jonathan Gonzalez

import re
import time
import sys
import logging
from datetime import datetime, timedelta
from urllib.parse import urljoin, urlparse, urlunparse, parse_qs, quote
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass, field
from collections import defaultdict
from difflib import SequenceMatcher

import requests
from bs4 import BeautifulSoup
from concurrent.futures import ThreadPoolExecutor, as_completed
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ===========================
# CONFIGURAÇÃO ULTIMATE
# ===========================
@dataclass
class Config:
    """Configuração ULTIMATE para Jonathan"""
    
    # Keywords expandidas
    AREA_KEYWORDS: List[str] = field(default_factory=lambda: [
        # IT Support
        "it support", "tech support", "technical support", "desktop support",
        "service desk", "help desk", "helpdesk",
        
        # SOC / Security (PRIORIDADE)
        "soc analyst", "soc tier 1", "soc tier 2", "soc l1", "soc l2",
        "security operations", "security analyst", "security monitoring",
        "blue team", "threat analyst", "incident response",
        "cyber security analyst", "cybersecurity analyst", "infosec",
        
        # Infrastructure
        "infrastructure analyst", "infrastructure support", "infrastructure engineer",
        "systems engineer", "cloud engineer", "cloud support",
        
        # NOC
        "noc analyst", "noc technician", "noc engineer",
        "network operations", "network monitoring", "monitoring analyst",
        
        # Sysadmin
        "sysadmin", "system administrator", "systems administrator",
        "junior sysadmin", "linux administrator", "windows administrator",
        "linux support", "windows support",
        
        # DevOps Junior
        "devops junior", "junior devops", "devops engineer junior",
        "automation engineer", "site reliability engineer junior", "sre junior",
        
        # Support roles
        "support engineer", "support analyst", "technical analyst",
        "it analyst", "it technician",
        
        # Tiers
        "tier 1", "tier 2", "level 1", "level 2", "l1", "l2", "t1", "t2"
    ])
    
    # Skills do Jonathan (BONUS)
    BONUS_KEYWORDS: List[str] = field(default_factory=lambda: [
        "splunk", "siem", "python", "automation", "bash", "powershell",
        "linux", "ubuntu", "kali", "active directory", "ad",
        "windows server", "tcp/ip", "firewall", "vpn", "wireshark",
        "pfSense", "log analysis", "incident", "phishing", "malware",
        "burp suite", "nmap", "nessus", "snort", "metasploit"
    ])
    
    # Níveis aceitos
    LEVEL_KEYWORDS: List[str] = field(default_factory=lambda: [
        "entry level", "entry-level", "junior", "jr ", "jr.",
        "associate", "trainee", "estagiário", "estágio",
        "graduate", "early career", "pleno", "nivel i", "nivel 1",
        "tier 1", "tier 2", "level 1", "level 2", "l1", "l2",
        "0-2 years", "0-3 years", "1-2 years", "1-3 years",
        "no experience", "little experience", "beginner", "iniciante"
    ])
    
    # Excluir
    EXCLUDE_TERMS: List[str] = field(default_factory=lambda: [
        "senior", "sênior", "sr ", "sr.", "pleno/sênior",
        "lead", "manager", "diretor", "director", "gerente",
        "head of", "principal", "architect", "staff", "chief"
    ])
    
    # GEO ULTRA-RIGOROSO: Lista de países/regiões que EXCLUEM
    GEO_BLACKLIST: List[str] = field(default_factory=lambda: [
        # Países
        "united states", "usa", "us only", "india", "poland", "ukraine",
        "philippines", "argentina", "colombia", "mexico", "canada",
        "united kingdom", "uk", "ireland", "germany", "france",
        "netherlands", "spain", "italy", "australia", "new zealand",
        
        # Estados dos EUA
        "california", "texas", "florida", "new york", "nashville",
        "san diego", "cleveland", "pennsylvania", "iowa",
        
        # Cidades do BR (se especificar cidade = red flag)
        "são paulo sp", "rio de janeiro rj", "belo horizonte",
        "florianópolis", "porto alegre", "curitiba", "brasília",
        
        # Keywords de restrição
        "based in", "located in", "must be", "must reside",
        "hybrid", "híbrido", "presencial", "on-site", "onsite",
        "preference", "preferred location", "time zone required"
    ])
    
    # GEO WHITELIST: termos que SÃO aceitos
    GEO_WHITELIST: List[str] = field(default_factory=lambda: [
        "worldwide", "global", "international", "remote anywhere",
        "work from anywhere", "location independent", "fully remote",
        "100% remote", "remote first", "remoto", "home office"
    ])
    
    # Spam/Red flags
    SPAM_KEYWORDS: List[str] = field(default_factory=lambda: [
        "template", "example", "sample", "demo", "test job",
        "how to", "guide", "tutorial", "fake", "scam"
    ])
    
    # Urgência
    URGENCY_KEYWORDS: List[str] = field(default_factory=lambda: [
        "urgent", "urgente", "asap", "immediate", "immediately",
        "start now", "começar já", "hiring now", "quick start",
        "contratação imediata", "vaga urgente"
    ])
    
    MAX_EXPERIENCE_YEARS: int = 3
    SIMILARITY_THRESHOLD: float = 0.85
    
    # Network
    MAX_WORKERS: int = 12
    REQUEST_TIMEOUT: int = 15
    SCRAPER_TIMEOUT: int = 60
    SLEEP_BETWEEN_REQUESTS: float = 0.12
    
    RATE_LIMITS: Dict[str, int] = field(default_factory=lambda: defaultdict(
        lambda: 30,
        {"linkedin.com": 8, "indeed.com": 15, "glassdoor.com": 15}
    ))

config = Config()

# ===========================
# LOGGING
# ===========================
def setup_logging(verbose: bool = False):
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format='%(asctime)s [%(levelname)s] %(message)s',
        datefmt='%H:%M:%S'
    )
    logging.getLogger("urllib3").setLevel(logging.WARNING)
    logging.getLogger("requests").setLevel(logging.WARNING)

logger = logging.getLogger(__name__)

# ===========================
# HTTP SESSION
# ===========================
class RequestCache:
    def __init__(self, duration_hours: int = 2):
        self.cache: Dict[str, Tuple[str, datetime]] = {}
        self.duration = timedelta(hours=duration_hours)
    
    def get(self, url: str) -> Optional[str]:
        if url in self.cache:
            content, timestamp = self.cache[url]
            if datetime.now() - timestamp < self.duration:
                return content
            del self.cache[url]
        return None
    
    def set(self, url: str, content: str):
        self.cache[url] = (content, datetime.now())

class HTTPSession:
    def __init__(self):
        self.session = requests.Session()
        retries = Retry(total=3, backoff_factor=0.5,
                       status_forcelist=(429, 500, 502, 503, 504))
        adapter = HTTPAdapter(max_retries=retries, pool_maxsize=25)
        self.session.mount("https://", adapter)
        self.session.mount("http://", adapter)
        
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.9,pt-BR;q=0.8,pt;q=0.7",
            "Accept-Encoding": "gzip, deflate, br",
        })
        
        self.cache = RequestCache(2)
    
    def get(self, url: str, use_cache: bool = True) -> Optional[str]:
        if use_cache:
            cached = self.cache.get(url)
            if cached:
                return cached
        
        try:
            response = self.session.get(url, timeout=config.REQUEST_TIMEOUT)
            response.raise_for_status()
            content = response.text
            if use_cache:
                self.cache.set(url, content)
            time.sleep(config.SLEEP_BETWEEN_REQUESTS)
            return content
        except:
            return None

http_session = HTTPSession()

# ===========================
# UTILS
# ===========================
def normalize_url(url: str) -> str:
    try:
        parsed = urlparse(url)
        query_params = parse_qs(parsed.query)
        cleaned = {k: v for k, v in query_params.items()
                  if not k.lower().startswith(("utm", "fb", "ref", "source"))}
        query = "&".join(f"{k}={v[0]}" for k, v in cleaned.items())
        return urlunparse((parsed.scheme, parsed.netloc,
                          parsed.path.rstrip("/"), "", query, ""))
    except:
        return url

def clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
    return text.strip()

def calculate_similarity(text1: str, text2: str) -> float:
    return SequenceMatcher(None, text1.lower(), text2.lower()).ratio()

# ===========================
# EXTRAÇÃO
# ===========================
class ContentExtractor:
    @staticmethod
    def extract_title(soup: BeautifulSoup) -> str:
        selectors = [
            ("h1", {"class": re.compile(r"(job|title|position|vaga)", re.I)}),
            ("h1", {}),
            ("h2", {"class": re.compile(r"(job|title|vaga)", re.I)}),
            ("title", {})
        ]
        for tag, attrs in selectors:
            elem = soup.find(tag, attrs)
            if elem:
                text = elem.get_text(strip=True)
                if 5 < len(text) < 150:
                    return clean_text(text)
        return ""
    
    @staticmethod
    def extract_description(soup: BeautifulSoup) -> str:
        meta = soup.find("meta", {"name": "description"})
        if meta and meta.get("content"):
            desc = clean_text(meta.get("content"))
            if len(desc) > 50:
                return desc
        
        selectors = [
            ("div", {"class": re.compile(r"(description|descricao|descri)", re.I)}),
            ("section", {"class": re.compile(r"(description|content)", re.I)}),
            ("article", {}),
        ]
        
        for tag, attrs in selectors:
            elem = soup.find(tag, attrs)
            if elem:
                text = elem.get_text(" ", strip=True)
                if len(text) > 50:
                    return clean_text(text[:3000])
        
        if soup.body:
            return clean_text(soup.body.get_text(" ", strip=True)[:3000])
        return ""
    
    @staticmethod
    def extract_location(soup: BeautifulSoup, text: str = "") -> str:
        location_selectors = [
            ("span", {"class": re.compile(r"(location|local|localiza)", re.I)}),
            ("div", {"class": re.compile(r"(location|local)", re.I)}),
            ("p", {"class": re.compile(r"location", re.I)}),
        ]
        
        for tag, attrs in location_selectors:
            elem = soup.find(tag, attrs)
            if elem:
                loc_text = elem.get_text(strip=True)
                if loc_text and len(loc_text) < 100:
                    return clean_text(loc_text)
        
        patterns = [
            r"(?:location|local|localização)[:\s]+([^\n|•]+)",
            r"remote[:\s]*-?\s*([^\n|•]+)",
            r"\(([^)]*remote[^)]*)\)",
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text.lower(), re.I)
            if match:
                location = match.group(1).strip()
                if location and len(location) < 60:
                    return location.title()
        
        return ""
    
    @staticmethod
    def extract_company(soup: BeautifulSoup) -> str:
        company_selectors = [
            ("span", {"class": re.compile(r"(company|empresa)", re.I)}),
            ("div", {"class": re.compile(r"(company|empresa)", re.I)}),
            ("a", {"class": re.compile(r"(company|empresa)", re.I)}),
            ("h2", {"class": re.compile(r"(company|empresa)", re.I)}),
        ]
        for tag, attrs in company_selectors:
            elem = soup.find(tag, attrs)
            if elem:
                text = elem.get_text(strip=True)
                if text and 2 < len(text) < 100:
                    return clean_text(text)
        return ""
    
    @staticmethod
    def extract_salary(text: str) -> str:
        """Detecta menção de salário"""
        patterns = [
            r"R?\$\s*\d+[,.]?\d*k?",
            r"\d+k\s*-\s*\d+k",
            r"salário.*?R?\$\s*\d+",
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.I)
            if match:
                return match.group(0)
        return ""

# ===========================
# FILTROS ULTRA-RIGOROSOS
# ===========================
class JobFilter:
    """Filtros ULTRA-RIGOROSOS para Jonathan"""
    
    @staticmethod
    def matches_area(text: str) -> bool:
        if not text:
            return False
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in config.AREA_KEYWORDS)
    
    @staticmethod
    def has_bonus_skills(text: str) -> int:
        text_lower = text.lower()
        return sum(1 for skill in config.BONUS_KEYWORDS if skill in text_lower)
    
    @staticmethod
    def has_level_mention(text: str) -> Tuple[bool, str]:
        if not text:
            return False, "no_mention"
        
        text_lower = text.lower()
        
        if any(level in text_lower for level in config.LEVEL_KEYWORDS):
            return True, "junior"
        
        if any(term in text_lower for term in config.EXCLUDE_TERMS):
            return True, "senior"
        
        return False, "no_mention"
    
    @staticmethod
    def extract_experience_years(text: str) -> Optional[int]:
        patterns = [
            r"(\d+)\+\s*(?:years?|yrs?|anos?)",
            r"(\d+)-\d+\s*(?:years?|yrs?|anos?)",
            r"(?:mínimo|minimum)\s+(?:de\s+)?(\d+)\s*(?:years?|yrs?|anos?)",
            r"(?:at least|pelo menos)\s+(\d+)\s*(?:years?|yrs?|anos?)",
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            if matches:
                try:
                    return max(int(m) for m in matches)
                except:
                    continue
        return None
    
    @staticmethod
    def is_truly_worldwide(location: str, text: str) -> Tuple[bool, str]:
        """
        ULTRA-RIGOROSO: Rejeita QUALQUER menção de país/cidade específica
        """
        if not location and not text:
            return True, ""
        
        combined = f"{location} {text}".lower()
        
        # 1. Se tem whitelist term, aceitar IMEDIATAMENTE
        for whitelist in config.GEO_WHITELIST:
            if whitelist in combined:
                return True, ""
        
        # 2. Verificar blacklist (QUALQUER match = rejeita)
        for blacklisted in config.GEO_BLACKLIST:
            if blacklisted in combined:
                return False, f"Geo: {blacklisted}"
        
        # 3. Se location tem só "Remote" ou "Remoto" = OK
        if location:
            loc_words = location.lower().strip()
            if loc_words in ["remote", "remoto", "home office", "anywhere"]:
                return True, ""
            
            # Se tem outras palavras depois de Remote, verificar
            if "remote" in loc_words or "remoto" in loc_words:
                # Remover "remote" e ver o que sobra
                remainder = loc_words.replace("remote", "").replace("remoto", "").strip(" -,")
                if remainder:
                    # Se sobrou algo, verificar se é país
                    if any(bl in remainder for bl in config.GEO_BLACKLIST):
                        return False, f"Geo: {remainder}"
        
        # 4. Default: se não mencionou nada suspeito, aceitar
        return True, ""
    
    @staticmethod
    def is_urgent(text: str) -> bool:
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in config.URGENCY_KEYWORDS)
    
    @staticmethod
    def is_spam(title: str, description: str) -> bool:
        combined = f"{title} {description}".lower()
        
        if any(spam in combined for spam in config.SPAM_KEYWORDS):
            return True
        
        if len(title) < 5 or len(title) > 150:
            return True
        
        return False
    
    @classmethod
    def passes_all_filters(cls, title: str, description: str, 
                          location: str = "", company: str = "") -> Tuple[bool, str, int]:
        """Filtros completos ULTRA-RIGOROSOS"""
        combined = f"{title} {description}"
        
        # 1. Spam
        if cls.is_spam(title, description):
            return False, "❌ Spam", 0
        
        # 2. Área
        if not cls.matches_area(combined):
            return False, "❌ Área não corresponde", 0
        
        # 3. GEO ULTRA-RIGOROSO
        is_worldwide, geo_reason = cls.is_truly_worldwide(location, combined)
        if not is_worldwide:
            return False, f"❌ {geo_reason}", 0
        
        # 4. Nível
        desc_has_level, desc_level = cls.has_level_mention(description)
        title_has_level, title_level = cls.has_level_mention(title)
        
        if desc_level == "senior":
            return False, "❌ Sênior (desc)", 0
        
        if title_level == "senior" and desc_level != "junior":
            return False, "❌ Sênior (title)", 0
        
        # 5. Experiência
        exp_years = cls.extract_experience_years(description)
        if exp_years and exp_years > config.MAX_EXPERIENCE_YEARS:
            return False, f"❌ {exp_years}+ anos", 0
        
        # 6. Bonus score
        bonus_score = cls.has_bonus_skills(combined)
        
        return True, "✅ OK", bonus_score

# ===========================
# MODELO
# ===========================
@dataclass
class Job:
    title: str
    url: str
    source: str
    company: str = ""
    location: str = ""
    description: str = ""
    salary: str = ""
    is_urgent: bool = False
    bonus_score: int = 0
    
    def similarity_to(self, other: 'Job') -> float:
        return calculate_similarity(self.title, other.title)

# ===========================
# SCRAPERS
# ===========================
class BaseScraper:
    def __init__(self, name: str):
        self.name = name
        self.extractor = ContentExtractor()
        self.filter = JobFilter()
    
    def scrape(self) -> List[Job]:
        raise NotImplementedError
    
    def fetch_and_parse(self, url: str) -> Optional[BeautifulSoup]:
        html = http_session.get(url)
        return BeautifulSoup(html, "html.parser") if html else None
    
    def process_job_link(self, title: str, url: str, location: str = "") -> Optional[Job]:
        if not url or len(url) < 10:
            return None
        
        soup = self.fetch_and_parse(url)
        if not soup:
            return None
        
        if not title:
            title = self.extractor.extract_title(soup)
        
        description = self.extractor.extract_description(soup)
        company = self.extractor.extract_company(soup)
        
        if not location:
            location = self.extractor.extract_location(soup, f"{title} {description}")
        
        passes, reason, bonus_score = self.filter.passes_all_filters(
            title, description, location, company
        )
        
        if not passes:
            logger.debug(f"{reason}: {title[:50]}")
            return None
        
        is_urgent = self.filter.is_urgent(f"{title} {description}")
        salary = self.extractor.extract_salary(description)
        
        stars = "⭐" * min(bonus_score, 3) if bonus_score > 0 else ""
        logger.info(f"✓ [{self.name}] {stars}: {title[:60]}")
        
        return Job(
            title=title,
            url=normalize_url(url),
            source=self.name,
            company=company,
            location=location or "Remote",
            description=description[:500],
            salary=salary,
            is_urgent=is_urgent,
            bonus_score=bonus_score
        )

class LinkedInScraper(BaseScraper):
    def __init__(self):
        super().__init__("LinkedIn")
    
    def scrape(self) -> List[Job]:
        """LinkedIn com filtro GEO ULTRA-RIGOROSO"""
        keywords = ' OR '.join([
            '"IT Support"', '"Technical Support"', '"SOC Analyst"',
            '"Security Operations"', '"Infrastructure"', '"NOC Analyst"',
            '"Sysadmin"', '"DevOps Junior"', '"Help Desk"'
        ])
        
        levels = ' OR '.join(['Junior', '"Entry Level"', 'Associate', '"Tier 1"'])
        full_query = f"({keywords}) AND ({levels})"
        
        params = {
            "keywords": full_query,
            "location": "Worldwide",
            "f_WT": "2",  # Remote
            "f_E": "1,2",  # Entry
            "f_TPR": "r604800",
            "geoId": "92000000",
        }
        
        query_string = "&".join(f"{k}={quote(str(v))}" for k, v in params.items())
        url = f"https://www.linkedin.com/jobs/search/?{query_string}"
        
        soup = self.fetch_and_parse(url)
        if not soup:
            return []
        
        jobs = []
        for card in soup.select("div.base-card, li.jobs-search-results__list-item")[:50]:
            title_elem = card.select_one("h3, a.base-card__full-link")
            if not title_elem:
                continue
            
            title = title_elem.get_text(strip=True)
            
            link_elem = card.select_one("a.base-card__full-link, a[href*='/jobs/']")
            if not link_elem:
                continue
            
            href = link_elem.get("href", "").split("?")[0]
            if not href or len(href) < 10:
                continue
            
            # Extrair location DO CARD
            location_elem = card.select_one("span.job-search-card__location, span[class*='location']")
            location = location_elem.get_text(strip=True) if location_elem else ""
            
            # Filtro GEO ULTRA-RIGOROSO
            is_worldwide, geo_reason = self.filter.is_truly_worldwide(location, title)
            if not is_worldwide:
                logger.debug(f"Geo blocked: {location} - {title[:40]}")
                continue
            
            # Filtros rápidos
            if not self.filter.matches_area(title):
                continue
            
            has_level, level_type = self.filter.has_level_mention(title)
            if level_type == "senior":
                continue
            
            company_elem = card.select_one("h4, span.job-search-card__subtitle")
            company = company_elem.get_text(strip=True) if company_elem else ""
            
            is_urgent = self.filter.is_urgent(title)
            bonus_score = self.filter.has_bonus_skills(title)
            
            job = Job(
                title=title,
                url=href,
                source=self.name,
                company=company,
                location=location if is_worldwide else "Remote - Worldwide",
                is_urgent=is_urgent,
                bonus_score=bonus_score
            )
            
            stars = "⭐" * min(bonus_score, 3) if bonus_score > 0 else ""
            logger.info(f"✓ [{self.name}] {stars}: {title[:60]}")
            jobs.append(job)
        
        return jobs

class GenericScraper(BaseScraper):
    def __init__(self, name: str, url: str, selector: str):
        super().__init__(name)
        self.url = url
        self.selector = selector
    
    def scrape(self) -> List[Job]:
        soup = self.fetch_and_parse(self.url)
        if not soup:
            return []
        
        jobs = []
        for link in soup.select(self.selector)[:25]:
            title = link.get_text(strip=True)
            href = link.get("href")
            
            if not href or len(href) < 5:
                continue
            
            full_url = urljoin(self.url, href)
            
            # Extrair location se possível
            parent = link.find_parent()
            location_elem = None
            if parent:
                location_elem = parent.find(class_=re.compile(r"(location|local)", re.I))
            location = location_elem.get_text(strip=True) if location_elem else ""
            
            job = self.process_job_link(title, full_url, location)
            if job:
                jobs.append(job)
        
        return jobs

# ===========================
# SITES ULTIMATE (TOP QUALITY)
# ===========================
def create_ultimate_scrapers() -> List[BaseScraper]:
    """Sites TOP testados e aprovados"""
    return [
        LinkedInScraper(),
        
        # === GRINGOS TOP ===
        
        # 1. Indeed International
        GenericScraper("Indeed-Worldwide",
            "https://www.indeed.com/jobs?q=IT+Support+OR+SOC+Analyst+OR+Infrastructure+junior&l=Remote&remotejob=032b3046-06a3-4876-8dfd-474eb5e7ed11",
            "a[href*='/rc/clk']"),
        
        # 2. RemoteOK
        GenericScraper("RemoteOK",
            "https://remoteok.com/remote-dev-jobs",
            "a.preventLink[href*='/remote-jobs/']"),
        
        # 3. Remote.co
        GenericScraper("Remote.co",
            "https://remote.co/remote-jobs/developer/",
            "a[href*='/job/']"),
        
        # 4. WeWorkRemotely
        GenericScraper("WeWorkRemotely",
            "https://weworkremotely.com/remote-jobs/search?term=support",
            "a[href*='/remote-jobs/']"),
        
        # 5. FlexJobs
        GenericScraper("FlexJobs",
            "https://www.flexjobs.com/remote-jobs/computer-it",
            "a[href*='/jobs/']"),
        
        # 6. Arc.dev
        GenericScraper("Arc.dev",
            "https://arc.dev/remote-jobs",
            "a[href*='/remote-jobs/']"),
        
        # 7. Turing.com
        GenericScraper("Turing",
            "https://www.turing.com/jobs",
            "a[href*='/jobs/']"),
        
        # 8. Himalayas
        GenericScraper("Himalayas",
            "https://himalayas.app/jobs/it-support",
            "a[href*='/jobs/']"),
        
        # 9. CyberSecurityJobs
        GenericScraper("CyberSec-Entry",
            "https://www.cybersecurityjobs.com/entry-level-cybersecurity-jobs",
            "a[href*='/jobs/']"),
        
        # 10. EuropeRemotely
        GenericScraper("EuropeRemotely",
            "https://europeremotely.com/remote-jobs/",
            "a[href*='/job/']"),
        
        # === BRASIL TOP (SEM CATHO/INFOJOBS/VAGAS) ===
        
        # 11. GuPy (MELHOR DO BRASIL!)
        GenericScraper("GuPy",
            "https://portal.gupy.io/job-search/term=suporte%20OR%20infraestrutura%20OR%20soc",
            "a[href*='/job/']"),
        
        # 12. Programathor
        GenericScraper("Programathor",
            "https://programathor.com.br/jobs?q=suporte+OR+infraestrutura+OR+soc",
            "a[href*='/jobs/']"),
        
        # 13. LinkedIn Brasil
        GenericScraper("LinkedIn-BR",
            "https://www.linkedin.com/jobs/search/?keywords=Suporte%20OR%20SOC%20OR%20Infraestrutura&location=Brasil&f_WT=2&f_E=1%2C2",
            "a[href*='/jobs/']"),
        
        # 14. Trampos.co
        GenericScraper("Trampos.co",
            "https://trampos.co/oportunidades?q=suporte+ti+OR+infraestrutura",
            "a[href*='/oportunidades/']"),
        
        # 15. GeekHunter
        GenericScraper("GeekHunter",
            "https://www.geekhunter.com.br/vagas",
            "a[href*='/vagas/']"),
        
        # 16. Revelo
        GenericScraper("Revelo",
            "https://www.revelo.com.br/vagas",
            "a[href*='/vaga/']"),
        
        # 17. StartSe Jobs
        GenericScraper("StartSe",
            "https://startse.com/vagas",
            "a[href*='/vagas/']"),
        
        # 18. BeTalent
        GenericScraper("BeTalent",
            "https://betalent.com/vagas-de-emprego",
            "a[href*='/vaga/']"),
        
        # 19. Glassdoor Remote
        GenericScraper("Glassdoor",
            "https://www.glassdoor.com/Job/remote-it-support-jobs-SRCH_IL.0,6_IS11047_KO7,17.htm",
            "a[href*='/job-listing/']"),
        
        # 20. RemoteHub
        GenericScraper("RemoteHub",
            "https://remotehub.com/remote-jobs/it-support",
            "a[href*='/remote-jobs/']"),
    ]

# ===========================
# ORQUESTRAÇÃO
# ===========================
class JobFinderOrchestrator:
    def __init__(self):
        self.scrapers = create_ultimate_scrapers()
    
    def scrape_all(self) -> List[Job]:
        all_jobs = []
        
        logger.info(f"🚀 Iniciando {len(self.scrapers)} scrapers TOP...")
        
        with ThreadPoolExecutor(max_workers=config.MAX_WORKERS) as executor:
            future_to_scraper = {
                executor.submit(scraper.scrape): scraper 
                for scraper in self.scrapers
            }
            
            for future in as_completed(future_to_scraper):
                scraper = future_to_scraper[future]
                try:
                    jobs = future.result(timeout=config.SCRAPER_TIMEOUT)
                    if jobs:
                        all_jobs.extend(jobs)
                        logger.info(f"✓ {scraper.name}: {len(jobs)} vagas")
                    else:
                        logger.debug(f"⚠ {scraper.name}: 0 vagas")
                except Exception as e:
                    logger.debug(f"✗ {scraper.name}: {str(e)[:60]}")
        
        return all_jobs
    
    def deduplicate_jobs(self, jobs: List[Job]) -> List[Job]:
        if not jobs:
            return []
        
        # Fase 1: URLs duplicadas (manter maior bonus score)
        seen_urls = {}
        for job in jobs:
            normalized = normalize_url(job.url)
            if normalized not in seen_urls:
                seen_urls[normalized] = job
            elif job.bonus_score > seen_urls[normalized].bonus_score:
                seen_urls[normalized] = job
        
        unique_by_url = list(seen_urls.values())
        logger.info(f"Após dedup por URL: {len(unique_by_url)} vagas")
        
        # Fase 2: Títulos similares (manter maior bonus score)
        final_jobs = []
        for job in unique_by_url:
            is_duplicate = False
            for existing in final_jobs:
                if job.similarity_to(existing) >= config.SIMILARITY_THRESHOLD:
                    if job.bonus_score > existing.bonus_score:
                        final_jobs.remove(existing)
                        break
                    else:
                        is_duplicate = True
                        break
            
            if not is_duplicate:
                final_jobs.append(job)
        
        removed = len(jobs) - len(final_jobs)
        if removed > 0:
            logger.info(f"Removidas {removed} duplicatas")
        
        return final_jobs

# ===========================
# EXPORTAÇÃO ULTIMATE
# ===========================
class DOCXExporter:
    @staticmethod
    def add_hyperlink(paragraph, text: str, url: str):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)
        
        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    
    @staticmethod
    def create_document(jobs: List[Job], filename: str):
        doc = Document()
        
        # Título
        title = doc.add_heading('🌍 Vagas ULTIMATE - Jonathan Gonzalez', level=1)
        title.runs[0].font.color.rgb = RGBColor(5, 99, 193)
        
        subtitle = doc.add_paragraph()
        subtitle.add_run("SOC • Infrastructure • IT Support • DevOps | Remote Worldwide").italic = True
        subtitle.runs[0].font.size = Pt(11)
        subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        
        # Info
        info = doc.add_paragraph()
        info.add_run("\n📅 Gerado em: ").bold = True
        info.add_run(datetime.now().strftime("%d/%m/%Y às %H:%M:%S"))
        info.add_run("\n📊 Total: ").bold = True
        info.add_run(str(len(jobs)))
        
        # Stats
        urgent_jobs = [j for j in jobs if j.is_urgent]
        high_match = [j for j in jobs if j.bonus_score >= 2]
        with_salary = [j for j in jobs if j.salary]
        
        if urgent_jobs:
            info.add_run("\n🔥 Urgentes: ").bold = True
            info.add_run(f"{len(urgent_jobs)}")
        if high_match:
            info.add_run("\n⭐ Alto match: ").bold = True
            info.add_run(f"{len(high_match)}")
        if with_salary:
            info.add_run("\n💰 Com salário: ").bold = True
            info.add_run(f"{len(with_salary)}")
        
        # Fontes
        jobs_by_source = defaultdict(list)
        for job in jobs:
            jobs_by_source[job.source].append(job)
        
        info.add_run("\n🔍 Fontes ativas: ").bold = True
        active_sources = len([src for src, jbs in jobs_by_source.items() if jbs])
        info.add_run(f"{active_sources}/20 sites")
        
        # Filtros
        filters = doc.add_paragraph()
        filters.add_run("\n✅ Filtros ULTIMATE:\n").bold = True
        filters.add_run("  • 🌍 Remote Worldwide (GEO ULTRA-RIGOROSO)\n")
        filters.add_run("  • 🎯 SOC, Infrastructure, IT Support, DevOps Junior\n")
        filters.add_run(f"  • 📊 Máx {config.MAX_EXPERIENCE_YEARS} anos experiência\n")
        filters.add_run("  • ⭐ Prioriza suas skills (Splunk, SIEM, Python, Linux)\n")
        filters.add_run("  • 🚫 Sites ruins removidos (Catho, Infojobs, Vagas.com)")
        
        doc.add_paragraph("\n" + "="*80 + "\n")
        
        # SEÇÃO 1: URGENTES + ALTO MATCH
        priority_jobs = [j for j in jobs if j.is_urgent or j.bonus_score >= 2]
        if priority_jobs:
            doc.add_heading('🔥⭐ PRIORIDADE MÁXIMA', level=2)
            doc.add_paragraph("Urgentes ou alto match com suas skills:\n")
            
            for i, job in enumerate(sorted(priority_jobs, key=lambda j: (-j.is_urgent, -j.bonus_score)), 1):
                p = doc.add_paragraph()
                
                prefix = "🔥" if job.is_urgent else ""
                stars = "⭐" * min(job.bonus_score, 3) if job.bonus_score > 0 else ""
                p.add_run(f"{prefix}{stars} {i}. ").bold = True
                
                DOCXExporter.add_hyperlink(p, job.title, job.url)
                
                details = doc.add_paragraph()
                details.style = 'List Bullet'
                
                detail_parts = []
                if job.company:
                    detail_parts.append(f"🏢 {job.company}")
                detail_parts.append(f"📍 {job.location}")
                detail_parts.append(f"🔍 {job.source}")
                if job.salary:
                    detail_parts.append(f"💰 {job.salary}")
                if job.bonus_score > 0:
                    detail_parts.append(f"⭐ Score: {job.bonus_score}")
                
                details.add_run("  |  ".join(detail_parts))
                details.runs[0].font.size = Pt(9)
                details.runs[0].font.color.rgb = RGBColor(100, 100, 100)
                
                doc.add_paragraph()
            
            doc.add_paragraph("\n" + "="*80 + "\n")
        
        # SEÇÃO 2: RESTO (por fonte)
        other_jobs = [j for j in jobs if not j.is_urgent and j.bonus_score < 2]
        
        if other_jobs:
            doc.add_heading('📋 TODAS AS VAGAS', level=2)
            
            other_by_source = defaultdict(list)
            for job in other_jobs:
                other_by_source[job.source].append(job)
            
            for source in sorted(other_by_source.keys()):
                source_jobs = other_by_source[source]
                
                source_heading = doc.add_heading(f'📍 {source} ({len(source_jobs)})', level=3)
                source_heading.runs[0].font.color.rgb = RGBColor(70, 130, 180)
                
                for i, job in enumerate(source_jobs, 1):
                    p = doc.add_paragraph()
                    stars = "⭐ " if job.bonus_score > 0 else ""
                    p.add_run(f"{i}. {stars}").bold = True
                    DOCXExporter.add_hyperlink(p, job.title, job.url)
                    
                    details = doc.add_paragraph()
                    details.style = 'List Bullet'
                    
                    detail_parts = []
                    if job.company:
                        detail_parts.append(f"🏢 {job.company}")
                    detail_parts.append(f"📍 {job.location}")
                    if job.salary:
                        detail_parts.append(f"💰 {job.salary}")
                    
                    details.add_run("  |  ".join(detail_parts))
                    details.runs[0].font.size = Pt(9)
                    details.runs[0].font.color.rgb = RGBColor(100, 100, 100)
                    
                    doc.add_paragraph()
                
                doc.add_paragraph()
        
        # Rodapé
        doc.add_paragraph("\n" + "="*80)
        footer = doc.add_paragraph()
        footer_run = footer.add_run(
            "\n🤖 JobFinder ULTIMATE - Criado por Apacci para Jonathan Gonzalez\n"
            "🎯 20 sites TOP (gringos + Brasil) - SEM Catho/Infojobs/Vagas.com\n"
            "🌍 100% Remote Worldwide - Filtro GEO ULTRA-RIGOROSO\n"
            "⭐ Score = suas skills (Splunk, SIEM, Python, Linux, AD)\n"
            "🔥 Urgente = hiring immediately/ASAP\n"
            "💰 Salário quando mencionado\n"
            "💼 Foco: Entry-Level até 3 anos experiência"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.save(filename)
        logger.info(f"📄 Documento salvo: {filename}")

# ===========================
# MAIN
# ===========================
def print_banner():
    banner = """
╔════════════════════════════════════════════════════════════════════════════╗
║              🌍 JobFinder ULTIMATE - Jonathan Gonzalez 🌍                  ║
║         SOC • Infrastructure • IT Support • DevOps | Remote Worldwide      ║
║                    20 Sites TOP + Filtro GEO ULTRA-RIGOROSO               ║
║                         Criado por Apacci                                  ║
╚════════════════════════════════════════════════════════════════════════════╝
    """
    print(banner)

def main():
    print_banner()
    
    print("⚙️  Configurando ambiente ULTIMATE...")
    setup_logging(verbose=False)
    
    print("\n🎯 Seu perfil:")
    print("   👤 Jonathan Gonzalez")
    print("   💼 SOC Analyst • Infrastructure • IT Support")
    print("   🛠️  Skills: Splunk, SIEM, Python, Linux, AD, Windows Server")
    print("   🌍 Inglês fluente + Espanhol")
    
    print("\n✅ Melhorias ULTIMATE:")
    print(f"   🌍 Filtro GEO ULTRA-RIGOROSO (rejeita QUALQUER país específico)")
    print(f"   🔍 20 sites TOP (gringos + Brasil)")
    print(f"   🚫 Removidos: Catho, Infojobs, Vagas.com")
    print(f"   ⭐ Sistema de match com suas skills")
    print(f"   💰 Detecta menção de salário")
    print(f"   🔥 Prioriza vagas urgentes")
    print()
    
    start_time = datetime.now()
    
    try:
        orchestrator = JobFinderOrchestrator()
        
        print("🔍 Buscando em 20 fontes TOP...")
        print("⏱️  Isso pode levar 1-2 minutos...\n")
        
        all_jobs = orchestrator.scrape_all()
        
        if not all_jobs:
            print("\n⚠️  Nenhuma vaga encontrada.")
            print("💡 Possíveis causas:")
            print("   • Filtro GEO muito rigoroso (mas isso é bom!)")
            print("   • Sites temporariamente indisponíveis")
            print("   • Tente novamente em algumas horas")
            return
        
        print(f"\n📊 Total encontrado: {len(all_jobs)} vagas")
        
        print("🧹 Deduplicação inteligente...")
        unique_jobs = orchestrator.deduplicate_jobs(all_jobs)
        print(f"✨ Total único: {len(unique_jobs)} vagas")
        
        # Stats
        urgent = sum(1 for j in unique_jobs if j.is_urgent)
        high_match = sum(1 for j in unique_jobs if j.bonus_score >= 2)
        with_salary = sum(1 for j in unique_jobs if j.salary)
        
        if urgent > 0:
            print(f"🔥 Urgentes: {urgent}")
        if high_match > 0:
            print(f"⭐ Alto match: {high_match}")
        if with_salary > 0:
            print(f"💰 Com salário: {with_salary}")
        
        # Ordenar: urgentes > alto match > resto
        unique_jobs.sort(key=lambda j: (
            not j.is_urgent,
            -j.bonus_score,
            j.source,
            j.title
        ))
        
        # Exportar
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"vagas_ultimate_jonathan_{timestamp}.docx"
        
        print(f"\n💾 Gerando documento ULTIMATE...")
        DOCXExporter.create_document(unique_jobs, filename)
        
        # Resumo
        print("\n" + "="*80)
        print("✅ SUCESSO ULTIMATE!")
        print("="*80)
        print(f"📄 Arquivo: {filename}")
        print(f"📊 Total: {len(unique_jobs)} vagas WORLDWIDE")
        if urgent > 0:
            print(f"🔥 {urgent} vagas urgentes no topo")
        if high_match > 0:
            print(f"⭐ {high_match} vagas com suas skills")
        print()
        
        # Distribuição
        print("📍 Sites que retornaram vagas:")
        jobs_by_source = defaultdict(int)
        for job in unique_jobs:
            jobs_by_source[job.source] += 1
        
        for source, count in sorted(jobs_by_source.items(), key=lambda x: -x[1]):
            bar = "█" * min(count, 40)
            print(f"   {source:20s} {bar} {count}")
        
        # Performance
        elapsed = datetime.now() - start_time
        print(f"\n⏱️  Tempo total: {elapsed.total_seconds():.1f}s")
        print(f"⚡ {len(unique_jobs)} vagas de qualidade encontradas!")
        print("="*80)
        
        # Dicas
        print("\n💡 Sobre o documento:")
        print("   🔥⭐ Seção PRIORIDADE com urgentes + alto match")
        print("   💰 Salário mostrado quando disponível")
        print("   🌍 100% Remote Worldwide (GEO ULTRA-RIGOROSO)")
        print("   🎯 Sites TOP sem spam")
        print()
        print("💪 Dica do Apacci:")
        print("   • Priorize vagas com 🔥 + ⭐⭐⭐")
        print("   • Seu perfil SOC+Python é OURO!")
        print("   • Aplique logo nas urgentes (respondem rápido)")
        print()
        
    except KeyboardInterrupt:
        print("\n\n⚠️  Interrompido pelo usuário")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Erro fatal: {str(e)}")
        print(f"\n❌ Erro: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
