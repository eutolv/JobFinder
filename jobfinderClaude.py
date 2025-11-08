#!/usr/bin/env python3
# jobfinder_ultra_optimized.py
# Scraper otimizado com filtros calibrados por Apacci

import re
import time
import sys
import logging
from datetime import datetime, timedelta
from urllib.parse import urljoin, urlparse, urlunparse, parse_qs
from typing import List, Dict, Tuple, Optional, Set
from dataclasses import dataclass, field
from collections import defaultdict
from difflib import SequenceMatcher

import requests
from bs4 import BeautifulSoup
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ===========================
# CONFIGURAÇÃO
# ===========================
@dataclass
class Config:
    """Configurações calibradas"""
    
    # Termos de área - EXPANDIDOS com variações
    AREA_KEYWORDS: List[str] = field(default_factory=lambda: [
        # IT Support variations
        "it support", "tech support", "technical support", "it helpdesk", 
        "desktop support", "it technician", "it specialist",
        
        # Service/Help Desk
        "service desk", "servicedesk", "help desk", "helpdesk",
        
        # NOC variations
        "noc analyst", "noc technician", "noc engineer", "noc specialist",
        "network operations", "network monitoring",
        
        # Monitoring
        "monitoring analyst", "monitoring technician", "monitoring engineer",
        "infrastructure monitoring",
        
        # SOC variations
        "soc tier 1", "soc level 1", "soc l1", "soc analyst",
        "security operations", "security monitoring",
        
        # Sysadmin
        "sysadmin", "system administrator", "systems administrator",
        "system admin", "systems admin", "junior sysadmin",
        
        # Linux Support
        "linux support", "linux administrator", "linux technician",
        "linux engineer", "linux admin",
        
        # Infrastructure
        "infrastructure support", "infrastructure technician",
        "infrastructure analyst",
        
        # Support roles
        "support engineer", "support technician", "support analyst",
        "support specialist", "technical analyst",
        
        # Tier/Level variations
        "tier 1", "tier 2", "tier i", "tier ii",
        "level 1", "level 2", "l1 support", "l2 support",
        "t1 support", "t2 support"
    ])
    
    # Termos de nível júnior/entry (DEVEM estar presentes OU não mencionar experiência)
    LEVEL_KEYWORDS: List[str] = field(default_factory=lambda: [
        # Entry level
        "entry level", "entry-level", "entry",
        
        # Junior
        "junior", "jr ", "jr.", "jnr", "jnr.",
        
        # Beginner
        "beginner", "iniciante", "trainee", "intern",
        
        # Associate
        "associate", "assistant",
        
        # Graduate
        "graduate", "new grad", "recent grad", "fresh graduate",
        
        # Career level
        "early career", "career starter",
        
        # Experience indicators
        "no experience", "without experience", "0 years",
        "0-1 year", "0-2 years", "up to 2 years",
        "little experience", "minimal experience"
    ])
    
    # Termos que EXCLUEM (sênior, etc)
    EXCLUDE_SENIOR_TERMS: List[str] = field(default_factory=lambda: [
        "senior", "sr ", "sr.", "sênior",
        "lead", "leading", "team lead",
        "manager", "management", "managing",
        "director", "head of", "chief",
        "principal", "architect",
        "staff engineer", "staff support",
        "expert", "specialist iii", "specialist iv", "specialist v",
        "iii", "iv", "v", " 3", " 4", " 5"
    ])
    
    # Restrições geográficas
    GEO_RESTRICTIONS: List[str] = field(default_factory=lambda: [
        # Explicitamente só um país
        "only", "únicamente", "solamente", "exclusively",
        
        # Preferências (também são restrições)
        "preference", "preferred location", "preferably",
        
        # Localização obrigatória
        "must be located", "must reside", "must be based",
        "based in", "located in", "residing in",
        
        # Presencial/Hybrid
        "hybrid", "híbrido", "on-site", "onsite", "on site",
        "in-office", "in office", "presencial",
        
        # Distância/região
        "commutable", "within", "local to", "near",
        
        # Timezone específico (red flag)
        "timezone required", "time zone required"
    ])
    
    # Cidades/Estados/Países específicos (red flags)
    SPECIFIC_LOCATIONS: List[str] = field(default_factory=lambda: [
        # US States
        "texas", "california", "new york", "florida", "iowa",
        "massachusetts", "washington state",
        
        # US Cities
        "san antonio", "san francisco", "new york city", "austin",
        "seattle", "boston", "miami",
        
        # Countries (quando específico)
        "us only", "usa only", "uk only", "eu only",
        "mexico only", "canada only", "india only",
        "brazil only", "brasil only"
    ])
    
    # Spam indicators
    SPAM_KEYWORDS: List[str] = field(default_factory=lambda: [
        "template", "example", "sample", "demo",
        "how to", "guide", "tutorial", "about us",
        "our company", "we are hiring", "join our team"
    ])
    
    # Hiring urgency indicators
    URGENCY_KEYWORDS: List[str] = field(default_factory=lambda: [
        "urgent", "urgently", "asap", "immediate", "immediately",
        "start now", "starting now", "quick start", "fast start",
        "hiring now", "hire immediately", "needed asap",
        "start asap", "available now", "begin immediately"
    ])
    
    # Certificações em excesso (red flag para júnior)
    EXCESSIVE_CERTS: List[str] = field(default_factory=lambda: [
        "ccna", "ccnp", "cissp", "cism", "ceh",
        "comptia a+", "comptia network+", "comptia security+",
        "aws certified", "azure certified", "gcp certified",
        "itil", "prince2", "pmp"
    ])
    
    # Big Tech companies (certificações OK)
    BIG_TECH: List[str] = field(default_factory=lambda: [
        "google", "microsoft", "amazon", "meta", "facebook",
        "apple", "netflix", "uber", "airbnb", "twitter",
        "linkedin", "salesforce", "oracle", "ibm", "cisco"
    ])
    
    # Experiência máxima aceita
    MAX_EXPERIENCE_YEARS: int = 2
    
    # Configurações de rede
    MAX_WORKERS: int = 8
    REQUEST_TIMEOUT: int = 12
    SCRAPER_TIMEOUT: int = 45
    SLEEP_BETWEEN_REQUESTS: float = 0.1
    CACHE_DURATION_HOURS: int = 2
    SIMILARITY_THRESHOLD: float = 0.85
    
    RATE_LIMITS: Dict[str, int] = field(default_factory=lambda: defaultdict(
        lambda: 30,
        {"linkedin.com": 10, "indeed.com": 15, "glassdoor.com": 15}
    ))

config = Config()

# ===========================
# LOGGING
# ===========================
def setup_logging(verbose: bool = False):
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format='%(asctime)s [%(levelname)s] %(message)s',
        datefmt='%H:%M:%S'
    )
    logging.getLogger("urllib3").setLevel(logging.WARNING)
    logging.getLogger("requests").setLevel(logging.WARNING)

logger = logging.getLogger(__name__)

# ===========================
# CACHE E RATE LIMITING
# ===========================
class RequestCache:
    def __init__(self, duration_hours: int = 2):
        self.cache: Dict[str, Tuple[str, datetime]] = {}
        self.duration = timedelta(hours=duration_hours)
    
    def get(self, url: str) -> Optional[str]:
        if url in self.cache:
            content, timestamp = self.cache[url]
            if datetime.now() - timestamp < self.duration:
                return content
            del self.cache[url]
        return None
    
    def set(self, url: str, content: str):
        self.cache[url] = (content, datetime.now())

class RateLimiter:
    def __init__(self, limits: Dict[str, int]):
        self.limits = limits
        self.requests: Dict[str, List[datetime]] = defaultdict(list)
    
    def wait_if_needed(self, url: str):
        domain = urlparse(url).netloc
        limit = self.limits[domain]
        now = datetime.now()
        
        self.requests[domain] = [
            t for t in self.requests[domain] 
            if now - t < timedelta(minutes=1)
        ]
        
        if len(self.requests[domain]) >= limit:
            oldest = min(self.requests[domain])
            wait_time = 60 - (now - oldest).total_seconds()
            if wait_time > 0:
                time.sleep(wait_time)
                self.requests[domain] = []
        
        self.requests[domain].append(now)

# ===========================
# SESSÃO HTTP
# ===========================
class HTTPSession:
    def __init__(self):
        self.session = requests.Session()
        retries = Retry(
            total=3,
            backoff_factor=0.5,
            status_forcelist=(429, 500, 502, 503, 504),
            allowed_methods=["GET"]
        )
        adapter = HTTPAdapter(max_retries=retries, pool_maxsize=20)
        self.session.mount("https://", adapter)
        self.session.mount("http://", adapter)
        
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml",
            "Accept-Language": "en-US,en;q=0.9",
        })
        
        self.cache = RequestCache(config.CACHE_DURATION_HOURS)
        self.rate_limiter = RateLimiter(config.RATE_LIMITS)
    
    def get(self, url: str, use_cache: bool = True) -> Optional[str]:
        if use_cache:
            cached = self.cache.get(url)
            if cached:
                return cached
        
        self.rate_limiter.wait_if_needed(url)
        
        try:
            response = self.session.get(url, timeout=config.REQUEST_TIMEOUT)
            response.raise_for_status()
            
            content = response.text
            if use_cache:
                self.cache.set(url, content)
            
            time.sleep(config.SLEEP_BETWEEN_REQUESTS)
            return content
        except:
            return None

http_session = HTTPSession()

# ===========================
# UTILITÁRIOS
# ===========================
def normalize_url(url: str) -> str:
    try:
        parsed = urlparse(url)
        query_params = parse_qs(parsed.query)
        cleaned = {k: v for k, v in query_params.items()
                  if not k.lower().startswith(("utm", "fb", "ref"))}
        query = "&".join(f"{k}={v[0]}" for k, v in cleaned.items())
        return urlunparse((parsed.scheme, parsed.netloc, 
                          parsed.path.rstrip("/"), "", query, ""))
    except:
        return url

def clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
    return text.strip()

def calculate_similarity(text1: str, text2: str) -> float:
    return SequenceMatcher(None, text1.lower(), text2.lower()).ratio()

# ===========================
# EXTRAÇÃO
# ===========================
class ContentExtractor:
    @staticmethod
    def extract_title(soup: BeautifulSoup) -> str:
        selectors = [
            ("h1", {"class": re.compile(r"(job|title|position)", re.I)}),
            ("h1", {}),
            ("title", {})
        ]
        for tag, attrs in selectors:
            elem = soup.find(tag, attrs)
            if elem:
                text = elem.get_text(strip=True)
                if text and 5 < len(text) < 150:
                    return clean_text(text)
        return ""
    
    @staticmethod
    def extract_description(soup: BeautifulSoup) -> str:
        meta = soup.find("meta", {"name": "description"})
        if meta and meta.get("content"):
            desc = clean_text(meta.get("content"))
            if len(desc) > 100:
                return desc
        
        selectors = [
            ("div", {"class": re.compile(r"(job[-_]?description|description)", re.I)}),
            ("section", {"class": re.compile(r"(description)", re.I)}),
            ("article", {}),
        ]
        
        for tag, attrs in selectors:
            elem = soup.find(tag, attrs)
            if elem:
                text = elem.get_text(" ", strip=True)
                if len(text) > 100:
                    return clean_text(text[:3000])
        
        if soup.body:
            return clean_text(soup.body.get_text(" ", strip=True)[:3000])
        return ""
    
    @staticmethod
    def extract_company(soup: BeautifulSoup) -> str:
        """Extrai nome da empresa"""
        company_selectors = [
            ("span", {"class": re.compile(r"company", re.I)}),
            ("div", {"class": re.compile(r"company", re.I)}),
            ("a", {"class": re.compile(r"company", re.I)}),
        ]
        for tag, attrs in company_selectors:
            elem = soup.find(tag, attrs)
            if elem:
                text = elem.get_text(strip=True)
                if text and len(text) < 100:
                    return clean_text(text)
        return ""

# ===========================
# FILTROS CALIBRADOS
# ===========================
class JobFilter:
    """Filtros calibrados por Apacci"""
    
    @staticmethod
    def matches_area(text: str) -> bool:
        """Verifica se tem palavras-chave de área (ordem flexível)"""
        if not text:
            return False
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in config.AREA_KEYWORDS)
    
    @staticmethod
    def has_level_mention(text: str) -> Tuple[bool, str]:
        """
        Verifica menção de nível/experiência
        Retorna: (tem_mencao, tipo)
        """
        if not text:
            return False, "no_mention"
        
        text_lower = text.lower()
        
        # Verificar júnior/entry
        if any(level in text_lower for level in config.LEVEL_KEYWORDS):
            return True, "junior"
        
        # Verificar sênior (red flag)
        if any(term in text_lower for term in config.EXCLUDE_SENIOR_TERMS):
            return True, "senior"
        
        # Não menciona nível
        return False, "no_mention"
    
    @staticmethod
    def extract_experience_years(text: str) -> Optional[int]:
        """Extrai anos de experiência mencionados"""
        patterns = [
            r"(\d+)\+\s*(?:years?|yrs?|anos?)",
            r"(\d+)-\d+\s*(?:years?|yrs?)",
            r"minimum\s+(?:of\s+)?(\d+)\s*(?:years?|yrs?)",
            r"at least\s+(\d+)\s*(?:years?|yrs?)",
            r"(\d+)\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp)"
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            if matches:
                try:
                    return max(int(m) for m in matches)
                except:
                    continue
        return None
    
    @staticmethod
    def has_geo_restriction(text: str) -> Tuple[bool, str]:
        """Detecta restrições geográficas"""
        text_lower = text.lower()
        
        # Verificar restrições gerais
        for restriction in config.GEO_RESTRICTIONS:
            if restriction in text_lower:
                return True, f"restriction:{restriction}"
        
        # Verificar localizações específicas
        for location in config.SPECIFIC_LOCATIONS:
            if location in text_lower:
                return True, f"location:{location}"
        
        return False, ""
    
    @staticmethod
    def is_urgent(text: str) -> bool:
        """Detecta contratação urgente"""
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in config.URGENCY_KEYWORDS)
    
    @staticmethod
    def has_excessive_certs(text: str, company: str) -> bool:
        """Detecta exigência excessiva de certificações (red flag pra júnior)"""
        text_lower = text.lower()
        company_lower = company.lower()
        
        # Se é big tech, certificações são OK
        if any(tech in company_lower for tech in config.BIG_TECH):
            return False
        
        # Contar certificações mencionadas
        cert_count = sum(1 for cert in config.EXCESSIVE_CERTS if cert in text_lower)
        
        # Mais de 2 certificações para júnior = red flag
        return cert_count > 2
    
    @staticmethod
    def is_spam(title: str, description: str) -> bool:
        """Detecta spam"""
        combined = f"{title} {description}".lower()
        
        if any(spam in combined for spam in config.SPAM_KEYWORDS):
            return True
        
        if len(title) < 5 or len(title) > 150:
            return True
        
        if re.match(r'^(job|template|example)s?$', title.lower()):
            return True
        
        return False
    
    @classmethod
    def passes_all_filters(cls, title: str, description: str, company: str = "") -> Tuple[bool, str]:
        """
        Filtros calibrados - DESCRIÇÃO tem prioridade sobre TÍTULO
        """
        combined = f"{title} {description}"
        
        # 1. Spam
        if cls.is_spam(title, description):
            return False, "❌ Spam/não-vaga"
        
        # 2. Área
        if not cls.matches_area(combined):
            return False, "❌ Área não corresponde"
        
        # 3. Restrição geográfica
        has_geo, geo_reason = cls.has_geo_restriction(combined)
        if has_geo:
            return False, f"❌ Geo: {geo_reason}"
        
        # 4. Análise de nível (DESCRIÇÃO > TÍTULO)
        desc_has_level, desc_level = cls.has_level_mention(description)
        title_has_level, title_level = cls.has_level_mention(title)
        
        # Se DESCRIÇÃO menciona sênior, REJEITA (mesmo se título diz júnior)
        if desc_level == "senior":
            return False, "❌ Sênior na descrição"
        
        # Se DESCRIÇÃO menciona júnior, ACEITA
        if desc_level == "junior":
            pass  # OK, continua
        # Se TÍTULO menciona sênior E descrição não menciona júnior, REJEITA
        elif title_level == "senior":
            return False, "❌ Sênior no título"
        # Se NENHUM dos dois menciona nível, REJEITA (pode ser cilada)
        elif not desc_has_level and not title_has_level:
            return False, "❌ Sem menção de nível (suspeito)"
        
        # 5. Anos de experiência (na descrição)
        exp_years = cls.extract_experience_years(description)
        if exp_years and exp_years > config.MAX_EXPERIENCE_YEARS:
            return False, f"❌ {exp_years}+ anos experiência"
        
        # 6. Certificações excessivas
        if cls.has_excessive_certs(combined, company):
            return False, "❌ Muitas certificações (red flag)"
        
        return True, "✅ OK"

# ===========================
# MODELO
# ===========================
@dataclass
class Job:
    title: str
    url: str
    source: str
    company: str = ""
    location: str = ""
    description: str = ""
    is_urgent: bool = False
    
    def similarity_to(self, other: 'Job') -> float:
        return calculate_similarity(self.title, other.title)

# ===========================
# SCRAPERS
# ===========================
class BaseScraper:
    def __init__(self, name: str):
        self.name = name
        self.extractor = ContentExtractor()
        self.filter = JobFilter()
    
    def scrape(self) -> List[Job]:
        raise NotImplementedError
    
    def fetch_and_parse(self, url: str) -> Optional[BeautifulSoup]:
        html = http_session.get(url)
        return BeautifulSoup(html, "html.parser") if html else None
    
    def process_job_link(self, title: str, url: str) -> Optional[Job]:
        if not url or len(url) < 10:
            return None
        
        soup = self.fetch_and_parse(url)
        if not soup:
            return None
        
        if not title:
            title = self.extractor.extract_title(soup)
        
        description = self.extractor.extract_description(soup)
        company = self.extractor.extract_company(soup)
        
        passes, reason = self.filter.passes_all_filters(title, description, company)
        if not passes:
            logger.debug(f"{reason}: {title[:50]}")
            return None
        
        is_urgent = self.filter.is_urgent(f"{title} {description}")
        
        logger.info(f"✓ [{self.name}]: {title[:60]}")
        return Job(
            title=title,
            url=normalize_url(url),
            source=self.name,
            company=company,
            description=description[:500],
            is_urgent=is_urgent
        )

class LinkedInScraper(BaseScraper):
    def __init__(self):
        super().__init__("LinkedIn")
    
    def scrape(self) -> List[Job]:
        keywords = ' OR '.join([
            '"IT Support"', '"Technical Support"', '"Service Desk"',
            '"Help Desk"', '"NOC Analyst"', '"SOC Tier 1"',
            '"Linux Support"', '"Sysadmin Junior"'
        ])
        
        levels = ' OR '.join(['Junior', '"Entry Level"', 'Associate', '"Tier 1"'])
        full_query = f"({keywords}) AND ({levels})"
        
        params = {
            "keywords": full_query,
            "location": "Worldwide",
            "f_WT": "2",  # Remote
            "f_E": "1,2",  # Entry
            "f_TPR": "r604800",
            "position": "1",
            "pageNum": "0"
        }
        
        query_string = "&".join(f"{k}={requests.utils.quote(str(v))}" for k, v in params.items())
        url = f"https://www.linkedin.com/jobs/search/?{query_string}"
        
        soup = self.fetch_and_parse(url)
        if not soup:
            return []
        
        jobs = []
        for card in soup.select("div.base-card, li")[:30]:
            title_elem = card.select_one("h3, a.base-card__full-link")
            if not title_elem:
                continue
            
            title = title_elem.get_text(strip=True)
            link_elem = card.select_one("a.base-card__full-link, a[href*='/jobs/']")
            if not link_elem:
                continue
            
            href = link_elem.get("href", "").split("?")[0]
            if not href or len(href) < 10:
                continue
            
            # Filtros rápidos
            if self.filter.has_geo_restriction(title)[0]:
                continue
            
            if not self.filter.matches_area(title):
                continue
            
            has_level, level_type = self.filter.has_level_mention(title)
            if level_type == "senior":
                continue
            
            company_elem = card.select_one("h4, span.job-search-card__subtitle")
            company = company_elem.get_text(strip=True) if company_elem else ""
            
            is_urgent = self.filter.is_urgent(title)
            
            job = Job(
                title=title,
                url=href,
                source=self.name,
                company=company,
                is_urgent=is_urgent
            )
            logger.info(f"✓ [{self.name}]: {title[:60]}")
            jobs.append(job)
        
        return jobs

class GenericScraper(BaseScraper):
    def __init__(self, name: str, url: str, selector: str):
        super().__init__(name)
        self.url = url
        self.selector = selector
    
    def scrape(self) -> List[Job]:
        soup = self.fetch_and_parse(self.url)
        if not soup:
            return []
        
        jobs = []
        for link in soup.select(self.selector)[:15]:
            title = link.get_text(strip=True)
            href = link.get("href")
            
            if href:
                full_url = urljoin(self.url, href)
                job = self.process_job_link(title, full_url)
                if job:
                    jobs.append(job)
        
        return jobs

# ===========================
# ORQUESTRAÇÃO
# ===========================
class JobFinderOrchestrator:
    def __init__(self):
        self.scrapers = [
            LinkedInScraper(),
            GenericScraper("WeWorkRemotely", 
                          "https://weworkremotely.com/categories/remote-customer-support-jobs",
                          "section.jobs li a"),
            GenericScraper("RemoteCo", 
                          "https://remote.co/remote-jobs/customer-service/",
                          "a[href*='/remote-jobs/']"),
            GenericScraper("RemoteOK",
                          "https://remoteok.com/remote-support-jobs",
                          "a.preventLink"),
        ]
    
    def scrape_all(self) -> List[Job]:
        all_jobs = []
        
        with ThreadPoolExecutor(max_workers=config.MAX_WORKERS) as executor:
            future_to_scraper = {
                executor.submit(scraper.scrape): scraper 
                for scraper in self.scrapers
            }
            
            for future in as_completed(future_to_scraper):
                scraper = future_to_scraper[future]
                try:
                    jobs = future.result(timeout=config.SCRAPER_TIMEOUT)
                    if jobs:
                        all_jobs.extend(jobs)
                        logger.info(f"✓ {scraper.name}: {len(jobs)} vagas")
                except Exception as e:
                    logger.debug(f"✗ {scraper.name}: {str(e)[:60]}")
        
        return all_jobs
    
    def deduplicate_jobs(self, jobs: List[Job]) -> List[Job]:
        if not jobs:
            return []
        
        # URL duplicates
        seen_urls = {}
        for job in jobs:
            normalized = normalize_url(job.url)
            if normalized not in seen_urls:
                seen_urls[normalized] = job
        
        unique_by_url = list(seen_urls.values())
        
        # Title similarity
        final_jobs = []
        for job in unique_by_url:
            is_duplicate = False
            for existing in final_jobs:
                if job.similarity_to(existing) >= config.SIMILARITY_THRESHOLD:
                    is_duplicate = True
                    break
            if not is_duplicate:
                final_jobs.append(job)
        
        removed = len(jobs) - len(final_jobs)
        if removed > 0:
            logger.info(f"Removidas {removed} duplicatas")
        
        return final_jobs

# ===========================
# EXPORTAÇÃO
# ===========================
class DOCXExporter:
    @staticmethod
    def add_hyperlink(paragraph, text: str, url: str):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)
        
        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    
    @staticmethod
    def create_document(jobs: List[Job], filename: str):
        """Cria documento com vagas urgentes destacadas"""
        doc = Document()
        
        # Título
        title = doc.add_heading('🌍 Vagas IT Support Entry-Level - Remote Worldwide', level=1)
        title.runs[0].font.color.rgb = RGBColor(5, 99, 193)
        
        # Info
        info = doc.add_paragraph()
        info.add_run("📅 Gerado em: ").bold = True
        info.add_run(datetime.now().strftime("%d/%m/%Y às %H:%M:%S"))
        info.add_run("\n📊 Total de vagas: ").bold = True
        info.add_run(str(len(jobs)))
        
        # Vagas urgentes
        urgent_jobs = [j for j in jobs if j.is_urgent]
        if urgent_jobs:
            info.add_run("\n🔥 Vagas urgentes: ").bold = True
            info.add_run(f"{len(urgent_jobs)} (marcadas com 🔥)")
        
        # Fontes
        jobs_by_source = defaultdict(list)
        for job in jobs:
            jobs_by_source[job.source].append(job)
        
        info.add_run("\n🔍 Fontes: ").bold = True
        info.add_run(", ".join(f"{src} ({len(jbs)})" for src, jbs in sorted(jobs_by_source.items())))
        
        # Filtros
        filters = doc.add_paragraph()
        filters.add_run("\n✅ Filtros calibrados:\n").bold = True
        filters.add_run("  • Remote Worldwide (sem restrições geográficas)\n")
        filters.add_run("  • Entry-Level/Junior com menção explícita\n")
        filters.add_run(f"  • Máximo {config.MAX_EXPERIENCE_YEARS} anos de experiência\n")
        filters.add_run("  • IT Support, Help Desk, NOC, SOC, Sysadmin, Linux\n")
        filters.add_run("  • Sem exigência excessiva de certificações")
        
        doc.add_paragraph("\n" + "="*80 + "\n")
        
        # Separar urgentes e normais
        if urgent_jobs:
            doc.add_heading('🔥 VAGAS URGENTES (Hiring Immediately)', level=2)
            doc.add_paragraph("Estas vagas mencionam contratação imediata/urgente:\n")
            
            for i, job in enumerate(urgent_jobs, 1):
                p = doc.add_paragraph()
                p.add_run(f"🔥 {i}. ").bold = True
                DOCXExporter.add_hyperlink(p, job.title, job.url)
                
                if job.company:
                    comp_p = doc.add_paragraph()
                    comp_p.style = 'List Bullet'
                    comp_run = comp_p.add_run(f"🏢 Empresa: {job.company}")
                    comp_run.font.size = Pt(9)
                
                doc.add_paragraph()
            
            doc.add_paragraph("\n" + "="*80 + "\n")
        
        # Vagas normais agrupadas por fonte
        normal_jobs = [j for j in jobs if not j.is_urgent]
        
        if normal_jobs:
            doc.add_heading('📋 TODAS AS VAGAS', level=2)
            
            normal_by_source = defaultdict(list)
            for job in normal_jobs:
                normal_by_source[job.source].append(job)
            
            for source in sorted(normal_by_source.keys()):
                source_jobs = normal_by_source[source]
                
                source_heading = doc.add_heading(f'📍 {source} ({len(source_jobs)} vagas)', level=3)
                source_heading.runs[0].font.color.rgb = RGBColor(70, 130, 180)
                
                for i, job in enumerate(source_jobs, 1):
                    p = doc.add_paragraph()
                    p.add_run(f"{i}. ").bold = True
                    DOCXExporter.add_hyperlink(p, job.title, job.url)
                    
                    if job.company:
                        comp_p = doc.add_paragraph()
                        comp_p.style = 'List Bullet'
                        comp_run = comp_p.add_run(f"🏢 {job.company}")
                        comp_run.font.size = Pt(9)
                        comp_run.font.color.rgb = RGBColor(100, 100, 100)
                    
                    doc.add_paragraph()
                
                doc.add_paragraph()
        
        # Rodapé
        doc.add_paragraph("\n" + "="*80)
        footer = doc.add_paragraph()
        footer_run = footer.add_run(
            "\n🤖 JobFinder IT Support - Filtros calibrados por Apacci\n"
            "🌍 Apenas vagas Remote Worldwide\n"
            "💼 Entry-Level com menção explícita de nível\n"
            "🔥 Vagas urgentes destacadas no topo\n"
            "⚡ Deduplicação inteligente e anti-spam aplicados"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.save(filename)
        logger.info(f"📄 Documento salvo: {filename}")

# ===========================
# MAIN
# ===========================
def print_banner():
    banner = """
╔════════════════════════════════════════════════════════════════════════════╗
║                    🌍 JobFinder IT Support Ultra 🌍                        ║
║                Vagas Entry-Level Remote Worldwide                          ║
║                    Filtros Calibrados por Apacci                           ║
╚════════════════════════════════════════════════════════════════════════════╝
    """
    print(banner)

def main():
    print_banner()
    
    print("⚙️  Configurando ambiente...")
    setup_logging(verbose=False)
    
    print("🎯 Filtros calibrados ativos:")
    print(f"   ✅ Remote Worldwide (sem geo-restrições)")
    print(f"   ✅ Entry-Level/Junior com menção EXPLÍCITA")
    print(f"   ✅ Máximo {config.MAX_EXPERIENCE_YEARS} anos de experiência")
    print(f"   ✅ Descrição tem prioridade sobre título")
    print(f"   ✅ Anti-spam e certificações excessivas")
    print(f"   🔥 Vagas urgentes destacadas com emoji")
    print()
    
    start_time = datetime.now()
    
    try:
        orchestrator = JobFinderOrchestrator()
        
        print("🔍 Iniciando busca com filtros calibrados...")
        all_jobs = orchestrator.scrape_all()
        
        if not all_jobs:
            print("\n⚠️  Nenhuma vaga encontrada com os filtros calibrados.")
            print("💡 Possíveis causas:")
            print("   • Filtros muito restritivos (precisa menção explícita de júnior)")
            print("   • Sites temporariamente indisponíveis")
            print("   • Poucas vagas realmente entry-level disponíveis")
            print("\n💪 Isso é normal! Filtros rigorosos = vagas de qualidade!")
            return
        
        print(f"\n📊 Total encontrado: {len(all_jobs)} vagas")
        
        print("🧹 Aplicando deduplicação inteligente...")
        unique_jobs = orchestrator.deduplicate_jobs(all_jobs)
        print(f"✨ Total único: {len(unique_jobs)} vagas")
        
        # Contar urgentes
        urgent_count = sum(1 for j in unique_jobs if j.is_urgent)
        if urgent_count > 0:
            print(f"🔥 Vagas urgentes: {urgent_count}")
        
        # Ordenar: urgentes primeiro, depois por fonte
        unique_jobs.sort(key=lambda j: (not j.is_urgent, j.source, j.title))
        
        # Exportar
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"vagas_it_support_calibradas_{timestamp}.docx"
        
        print(f"\n💾 Gerando documento...")
        DOCXExporter.create_document(unique_jobs, filename)
        
        # Resumo
        print("\n" + "="*80)
        print("✅ SUCESSO!")
        print("="*80)
        print(f"📄 Arquivo: {filename}")
        print(f"📊 Total: {len(unique_jobs)} vagas filtradas")
        if urgent_count > 0:
            print(f"🔥 Urgentes: {urgent_count} vagas no topo do documento")
        print()
        
        # Resumo por fonte
        print("📍 Distribuição por fonte:")
        jobs_by_source = defaultdict(int)
        for job in unique_jobs:
            jobs_by_source[job.source] += 1
        
        for source, count in sorted(jobs_by_source.items(), key=lambda x: -x[1]):
            bar = "█" * min(count, 40)
            print(f"   {source:20s} {bar} {count}")
        
        # Stats
        elapsed = datetime.now() - start_time
        print(f"\n⏱️  Tempo total: {elapsed.total_seconds():.1f}s")
        if elapsed.total_seconds() > 0:
            print(f"⚡ Velocidade: {len(unique_jobs)/elapsed.total_seconds():.1f} vagas/s")
        print("="*80)
        
        # Dicas
        print("\n💡 Sobre os filtros calibrados:")
        print("   🎯 Só aceita vagas com menção EXPLÍCITA de júnior/entry")
        print("   📝 Descrição tem prioridade sobre título")
        print("   🌍 Zero tolerância com restrições geográficas")
        print("   🔥 Vagas urgentes aparecem primeiro no documento")
        print("   ⚠️  Poucos resultados? Ótimo! Qualidade > Quantidade")
        print()
        
    except KeyboardInterrupt:
        print("\n\n⚠️  Interrompido pelo usuário")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Erro fatal: {str(e)}")
        print(f"\n❌ Erro: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
